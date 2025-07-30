"""
Policy Generator Service
Handles AI-powered policy generation and management
"""
import asyncio
import uuid
from datetime import datetime
from typing import List, Dict, Optional, Any

from models.policy_models import (
    PolicyProject, 
    PolicyGenerationRequest, 
    PolicyGenerationResponse,
    PolicyProjectStatus,
    GeneratedPolicy,
    PolicyExportRequest,
    PolicyExportResponse
)
from database.connection import get_database
from services.grc.knowledge_base import GRCKnowledgeBase
from services.grc.llm_manager import LLMManager
from utils.exceptions import ServiceError
import logging

logger = logging.getLogger(__name__)

class PolicyGeneratorService:
    """Service for managing policy generation projects"""
    
    def __init__(self):
        self.db = None
        
        # Initialize services with error handling
        try:
            self.grc_knowledge = GRCKnowledgeBase()
            logger.info("GRC Knowledge Base initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize GRC Knowledge Base: {e}")
            self.grc_knowledge = None
            
        try:
            self.llm_manager = LLMManager()
            logger.info("LLM Manager initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize LLM Manager: {e}")
            self.llm_manager = None
    
    def _ensure_db_connection(self):
        """Ensure database connection is available with retry mechanism"""
        if self.db is None:
            try:
                self.db = get_database()
                if self.db is None:
                    raise ServiceError("Database connection not available. Please check MongoDB connection.")
            except Exception as e:
                logger.error(f"Database connection failed: {e}")
                raise ServiceError(f"Database connection failed: {str(e)}")
        
        # Test the connection by trying a simple operation
        try:
            # Ping the database to ensure it's responsive
            self.db.command('ping')
        except Exception as e:
            logger.error(f"Database ping failed: {e}")
            # Reset connection and try once more
            self.db = None
            try:
                self.db = get_database()
                if self.db is None:
                    raise ServiceError("Database reconnection failed")
                self.db.command('ping')
            except Exception as retry_error:
                logger.error(f"Database reconnection failed: {retry_error}")
                raise ServiceError(f"Database is not responsive: {str(retry_error)}")
    
    async def create_policy_project(
        self, 
        request: PolicyGenerationRequest, 
        user_id: str
    ) -> PolicyGenerationResponse:
        """Create a new policy project and start generation"""
        try:
            self._ensure_db_connection()
            
            # Create policy project
            project_id = str(uuid.uuid4())
            
            project = PolicyProject(
                id=project_id,
                title=request.title,
                description=request.description,
                framework=request.framework,
                prompt=request.prompt,
                status=PolicyProjectStatus.GENERATING,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
                user_id=user_id
            )
            
            # Save to database
            projects_collection = self.db.policy_projects
            project_dict = project.dict()
            await projects_collection.insert_one(project_dict)
            
            # Start async policy generation
            asyncio.create_task(self._generate_policy_async(project_id, request))
            
            return PolicyGenerationResponse(
                project_id=project_id,
                status="started",
                message="Policy generation started successfully"
            )
            
        except Exception as e:
            logger.error(f"Failed to create policy project: {str(e)}")
            raise ServiceError(f"Failed to create policy project: {str(e)}")
    
    async def _generate_policy_async(self, project_id: str, request: PolicyGenerationRequest):
        """Generate policy content asynchronously"""
        try:
            self._ensure_db_connection()
            projects_collection = self.db.policy_projects
            
            # Update status to generating
            await projects_collection.update_one(
                {"id": project_id},
                {
                    "$set": {
                        "status": PolicyProjectStatus.GENERATING.value,
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            
            # Generate policy content using LLM
            if self.llm_manager:
                policy_content = await self._generate_policy_content(request)
            else:
                # Fallback to mock generation
                policy_content = self._generate_mock_policy(request)
            
            # Create generated policy
            generated_policy = GeneratedPolicy(
                id=str(uuid.uuid4()),
                content=policy_content,
                word_count=len(policy_content.split()),
                generated_at=datetime.utcnow()
            )
            
            # Update project with generated policy
            await projects_collection.update_one(
                {"id": project_id},
                {
                    "$set": {
                        "generated_policy": generated_policy.dict(),
                        "status": PolicyProjectStatus.COMPLETED.value,
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"Policy generation completed for project {project_id}")
            
        except Exception as e:
            logger.error(f"Policy generation failed for project {project_id}: {str(e)}")
            
            # Update project status to failed
            try:
                await projects_collection.update_one(
                    {"id": project_id},
                    {
                        "$set": {
                            "status": PolicyProjectStatus.FAILED.value,
                            "error_message": str(e),
                            "updated_at": datetime.utcnow()
                        }
                    }
                )
            except Exception as update_error:
                logger.error(f"Failed to update project status: {update_error}")
    
    async def _generate_policy_content(self, request: PolicyGenerationRequest) -> str:
        """Generate policy content using AI/LLM"""
        try:
            # Use the LLM manager to generate policy content
            prompt = self._build_policy_prompt(request)
            
            # Generate using LLM
            response = await self.llm_manager.generate_response(prompt=prompt)
            
            if response and len(response.strip()) > 100:  # Ensure we got a substantial response
                logger.info(f"Successfully generated policy content using LLM: {len(response)} characters")
                return response.strip()
            else:
                logger.warning("LLM response was too short, falling back to mock generation")
                return self._generate_mock_policy(request)
            
        except Exception as e:
            logger.error(f"LLM policy generation failed: {str(e)}")
            # Fallback to mock generation
            logger.info("Falling back to mock policy generation")
            return self._generate_mock_policy(request)
    
    def _build_policy_prompt(self, request: PolicyGenerationRequest) -> str:
        """Build the prompt for AI policy generation"""
        framework_details = {
            "ISO27001": "ISO 27001 Information Security Management Systems",
            "SOC2": "SOC 2 Service Organization Control 2",
            "NIST_CSF": "NIST Cybersecurity Framework",
            "PCI_DSS": "Payment Card Industry Data Security Standard",
            "GDPR": "General Data Protection Regulation",
            "HIPAA": "Health Insurance Portability and Accountability Act",
            "CUSTOM": "Custom Business Policy Framework"
        }
        
        framework_name = framework_details.get(request.framework, request.framework)
        
        return f"""You are an expert compliance policy writer with deep knowledge of {framework_name}. 

Create a comprehensive, professional compliance policy document with the following specifications:

**Policy Title:** {request.title}
**Compliance Framework:** {framework_name}
**Specific Requirements:** {request.prompt}

**Instructions:**
1. Write a complete, professional policy document that addresses all the specified requirements
2. Ensure full alignment with {framework_name} standards and controls
3. Include specific framework citations where appropriate
4. Use clear, authoritative language suitable for corporate governance
5. Structure the policy with proper sections and subsections
6. Include implementation guidance and measurable requirements

**Required Policy Structure:**
1. **Purpose and Scope** - Define the policy's objective and applicability
2. **Policy Statement** - Clear commitment and high-level requirements
3. **Roles and Responsibilities** - Define who does what
4. **Implementation Procedures** - Specific steps and requirements
5. **Monitoring and Compliance** - How compliance will be measured and monitored
6. **Enforcement** - Consequences of non-compliance
7. **Policy Maintenance** - Review and update procedures

**Formatting Requirements:**
- Use markdown formatting with clear headers (# ## ###)
- Include bullet points and numbered lists where appropriate
- Add **Framework Alignment:** notes in relevant sections
- Include effective date, review date, and version information
- Ensure the document is 1500-3000 words for comprehensive coverage

**Framework-Specific Requirements:**
- Reference specific {framework_name} controls and requirements
- Include compliance metrics and KPIs relevant to {framework_name}
- Address audit and assessment requirements
- Include risk management considerations

Generate a complete, ready-to-implement policy document that meets enterprise standards and regulatory requirements."""
    
    def _generate_mock_policy(self, request: PolicyGenerationRequest) -> str:
        """Generate mock policy content for fallback"""
        return f"""# {request.title}

## 1. PURPOSE AND SCOPE

This policy establishes guidelines and procedures for {request.prompt.lower()} in accordance with {request.framework} requirements. This policy applies to all employees, contractors, and third parties who have access to organizational resources.

**Framework Alignment:** This section satisfies {request.framework} control requirements for policy documentation and scope definition.

## 2. POLICY STATEMENT

Our organization is committed to maintaining the highest standards of {request.prompt.lower()} through:

- Implementation of appropriate controls and safeguards
- Regular monitoring and assessment of compliance
- Continuous improvement of our security posture
- Training and awareness programs for all personnel

**Framework Alignment:** This section addresses {request.framework} policy statement requirements.

## 3. ROLES AND RESPONSIBILITIES

### 3.1 Management
- Provide leadership and resources for policy implementation
- Ensure compliance with regulatory requirements
- Review and approve policy updates annually

### 3.2 IT Security Team
- Implement technical controls and monitoring systems
- Conduct regular security assessments
- Respond to security incidents and breaches

### 3.3 All Employees
- Comply with policy requirements and procedures
- Report security incidents promptly
- Participate in required training programs

**Framework Alignment:** This section satisfies {request.framework} requirements for role-based responsibilities.

## 4. IMPLEMENTATION PROCEDURES

### 4.1 Control Implementation
All controls specified in this policy shall be implemented according to {request.framework} guidelines:

1. **Risk Assessment**: Regular assessment of risks related to {request.prompt.lower()}
2. **Control Selection**: Implementation of appropriate controls based on risk analysis
3. **Monitoring**: Continuous monitoring of control effectiveness
4. **Review**: Regular review and update of controls as needed

### 4.2 Documentation Requirements
- All procedures must be documented and maintained
- Evidence of compliance must be collected and retained
- Regular audits must be conducted to verify effectiveness

**Framework Alignment:** This section addresses {request.framework} implementation and documentation requirements.

## 5. MONITORING AND COMPLIANCE

### 5.1 Performance Metrics
Key performance indicators for this policy include:
- Compliance assessment scores
- Number of incidents or violations
- Training completion rates
- Control implementation status

### 5.2 Audit and Review
- Annual policy review and update process
- Regular internal audits of policy compliance
- External audit preparation and support
- Corrective action tracking and resolution

**Framework Alignment:** This section satisfies {request.framework} monitoring and audit requirements.

## 6. ENFORCEMENT

Non-compliance with this policy may result in disciplinary action up to and including termination of employment or contract. All violations will be investigated and appropriate corrective measures will be taken.

## 7. POLICY MAINTENANCE

This policy will be reviewed annually or as required by changes in:
- Regulatory requirements
- Business operations
- Technology infrastructure
- Risk environment

**Effective Date:** {datetime.now().strftime('%Y-%m-%d')}
**Review Date:** {datetime.now().replace(year=datetime.now().year + 1).strftime('%Y-%m-%d')}
**Version:** 1.0

---
*This policy was generated by CompliAI Policy Generator in accordance with {request.framework} requirements.*"""
    
    async def list_policy_projects(self, user_id: str) -> List[PolicyProject]:
        """Get all policy projects for a user"""
        try:
            self._ensure_db_connection()
            projects_collection = self.db.policy_projects
            
            projects_cursor = projects_collection.find({"user_id": user_id})
            projects = []
            
            async for project_dict in projects_cursor:
                # Convert ObjectId to string if present
                if "_id" in project_dict:
                    del project_dict["_id"]
                
                project = PolicyProject(**project_dict)
                projects.append(project)
            
            return sorted(projects, key=lambda x: x.created_at, reverse=True)
            
        except Exception as e:
            logger.error(f"Failed to list policy projects: {str(e)}")
            raise ServiceError(f"Failed to retrieve policy projects: {str(e)}")
    
    async def get_policy_project(self, project_id: str, user_id: str) -> Optional[PolicyProject]:
        """Get a specific policy project"""
        try:
            self._ensure_db_connection()
            projects_collection = self.db.policy_projects
            
            project_dict = await projects_collection.find_one({
                "id": project_id,
                "user_id": user_id
            })
            
            if not project_dict:
                return None
            
            # Convert ObjectId to string if present
            if "_id" in project_dict:
                del project_dict["_id"]
            
            return PolicyProject(**project_dict)
            
        except Exception as e:
            logger.error(f"Failed to get policy project: {str(e)}")
            raise ServiceError(f"Failed to retrieve policy project: {str(e)}")
    
    async def update_policy_content(
        self, 
        project_id: str, 
        user_id: str, 
        content: str
    ) -> PolicyProject:
        """Update policy content"""
        try:
            self._ensure_db_connection()
            projects_collection = self.db.policy_projects
            
            # Update the generated policy content
            updated_policy = GeneratedPolicy(
                id=str(uuid.uuid4()),
                content=content,
                word_count=len(content.split()),
                generated_at=datetime.utcnow()
            )
            
            result = await projects_collection.update_one(
                {"id": project_id, "user_id": user_id},
                {
                    "$set": {
                        "generated_policy": updated_policy.dict(),
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            
            if result.matched_count == 0:
                raise ServiceError("Policy project not found")
            
            return await self.get_policy_project(project_id, user_id)
            
        except Exception as e:
            logger.error(f"Failed to update policy content: {str(e)}")
            raise ServiceError(f"Failed to update policy content: {str(e)}")
    
    async def export_policy(
        self, 
        request: PolicyExportRequest, 
        user_id: str
    ) -> bytes:
        """Export policy to different formats and return binary data"""
        try:
            project = await self.get_policy_project(request.project_id, user_id)
            
            if not project or not project.generated_policy:
                raise ServiceError("Policy project not found or not generated")
            
            # Generate export data based on format
            if request.format == "pdf":
                export_data = await self._export_to_pdf(project, request)
            elif request.format == "docx":
                export_data = await self._export_to_docx(project, request)
            else:  # txt/markdown
                export_data = project.generated_policy.content.encode('utf-8')
            
            return export_data
            
        except Exception as e:
            logger.error(f"Failed to export policy: {str(e)}")
            raise ServiceError(f"Failed to export policy: {str(e)}")
    
    async def _export_to_pdf(self, project: PolicyProject, request: PolicyExportRequest) -> bytes:
        """Export policy to PDF format"""
        try:
            import markdown2
            from weasyprint import HTML, CSS
            from io import BytesIO
            
            # Convert markdown to HTML
            html_content = markdown2.markdown(
                project.generated_policy.content,
                extras=['fenced-code-blocks', 'tables', 'header-ids']
            )
            
            # Create a complete HTML document with styling
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{project.title}</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    h1 {{
                        color: #2c3e50;
                        border-bottom: 3px solid #3498db;
                        padding-bottom: 10px;
                        margin-top: 30px;
                    }}
                    h2 {{
                        color: #34495e;
                        border-bottom: 2px solid #bdc3c7;
                        padding-bottom: 5px;
                        margin-top: 25px;
                    }}
                    h3 {{
                        color: #34495e;
                        margin-top: 20px;
                    }}
                    strong {{
                        color: #2c3e50;
                    }}
                    ul, ol {{
                        margin: 10px 0;
                        padding-left: 25px;
                    }}
                    li {{
                        margin: 5px 0;
                    }}
                    .header {{
                        text-align: center;
                        margin-bottom: 30px;
                        padding: 20px;
                        background-color: #f8f9fa;
                        border-radius: 8px;
                    }}
                    .metadata {{
                        font-size: 12px;
                        color: #666;
                        margin-top: 5px;
                    }}
                    .framework-alignment {{
                        background-color: #e3f2fd;
                        border-left: 4px solid #2196f3;
                        padding: 10px;
                        margin: 15px 0;
                        font-style: italic;
                        color: #1565c0;
                    }}
                    @page {{
                        margin: 2cm;
                        @bottom-center {{
                            content: "Page " counter(page) " of " counter(pages);
                            font-size: 10px;
                            color: #666;
                        }}
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; border: none;">{project.title}</h1>
                    <div class="metadata">
                        <strong>Framework:</strong> {project.framework} | 
                        <strong>Generated:</strong> {project.created_at.strftime('%B %d, %Y')} | 
                        <strong>Words:</strong> {project.generated_policy.word_count}
                    </div>
                </div>
                {html_content}
            </body>
            </html>
            """
            
            # Convert HTML to PDF
            pdf_buffer = BytesIO()
            HTML(string=full_html).write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
            
            return pdf_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"PDF export failed: {str(e)}")
            # Fallback: return the markdown content as bytes
            return project.generated_policy.content.encode('utf-8')
    
    async def _export_to_docx(self, project: PolicyProject, request: PolicyExportRequest) -> bytes:
        """Export policy to DOCX format"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            from io import BytesIO
            import re
            
            # Create a new document
            doc = Document()
            
            # Set up styles
            styles = doc.styles
            
            # Title style
            title_style = styles['Title']
            title_style.font.size = Pt(18)
            title_style.font.name = 'Calibri'
            
            # Add document header
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title_run = header_para.add_run(project.title)
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            
            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metadata_text = f"Framework: {project.framework} | Generated: {project.created_at.strftime('%B %d, %Y')} | Words: {project.generated_policy.word_count}"
            metadata_run = metadata_para.add_run(metadata_text)
            metadata_run.font.size = Pt(10)
            metadata_run.font.italic = True
            
            doc.add_paragraph()  # Add space
            
            # Process the markdown content
            lines = project.generated_policy.content.split('\n')
            current_list = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Handle headers
                if line.startswith('# '):
                    para = doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    para = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    para = doc.add_heading(line[4:], level=3)
                
                # Handle Framework Alignment boxes
                elif line.startswith('**Framework Alignment:**'):
                    para = doc.add_paragraph()
                    para.style = 'Intense Quote'
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    run = para.add_run(text)
                    run.font.size = Pt(10)
                    run.font.italic = True
                
                # Handle bullet points
                elif line.startswith('- '):
                    para = doc.add_paragraph(line[2:], style='List Bullet')
                
                # Handle numbered lists
                elif re.match(r'^\d+\. ', line):
                    para = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
                
                # Handle bold text and regular paragraphs
                else:
                    para = doc.add_paragraph()
                    
                    # Process bold text
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            run.font.bold = True
                        else:
                            para.add_run(part)
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"DOCX export failed: {str(e)}")
            # Fallback: return the markdown content as bytes
            return project.generated_policy.content.encode('utf-8')

# Global instance
policy_generator_service = PolicyGeneratorService()
